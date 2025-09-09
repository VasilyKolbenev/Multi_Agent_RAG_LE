import os, uuid, json
from typing import List, Optional
from fastapi import FastAPI, UploadFile, File, Form, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from sse_starlette.sse import EventSourceResponse
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

from retrieval import HybridCorpus, Document, EmbeddedDocument, update_document_in_corpus, get_corpus_stats, clear_corpus
from agents import MultiAgent
from llm import LLM
from graph_index import GraphIndex
from agentic_rag import AgenticRAGSystem
from storage import append_trace, read_traces
from langx import run_extraction, stream_extraction
from profiles import PROFILES

app = FastAPI(title="MultiAgent-RAG Pro")

# Настройка CORS для конкретных доменов
origins = [
    "https://vasilykolbenev.github.io",  # Ваш GitHub Pages
    "http://localhost:8001",           # Локальный сервер для тестов
    "http://localhost",                # На всякий случай
    "null",                            # Для локальных файлов (file://)
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)

corpus = HybridCorpus()
graph = GraphIndex()
agent = MultiAgent(corpus, graph)
agentic_system = AgenticRAGSystem(corpus, graph)

@app.get("/health")
def health(): return {"ok": True, "model": os.getenv("LLM_MODEL","gpt-5-mini")}

@app.get("/profiles")
def profiles(): return PROFILES

@app.get("/traces")
def traces(): return read_traces()

@app.post("/ingest/text")
async def ingest_text(text: str = Form(...), doc_id: Optional[str] = Form(None)):
    """Принимает текст, генерирует ID, если он не предоставлен."""
    try:
        if not doc_id:
            # Генерируем простой doc_id на основе хэша текста
            import hashlib
            doc_id = f"doc_{hashlib.md5(text[:100].encode()).hexdigest()[:8]}"
            
        corpus.ingest_text(doc_id, text)
        append_trace({"type":"ingest", "doc_id":doc_id, "len":len(text)})
        return {"ok": True, "doc_id": doc_id}
    except Exception as e:
        print(f"ERROR in /ingest/text: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(status_code=500, content={"message": "Internal Server Error during ingestion"})

@app.post("/ingest/files")
async def ingest_files(files: List[UploadFile] = File(...)):
    saved = []
    base = os.environ.get("DOCS_DIR", "data/docs"); os.makedirs(base, exist_ok=True)
    for f in files:
        fp = os.path.join(base, f.filename)
        with open(fp, "wb") as out: out.write(await f.read())
        saved.append(f.filename)
    
    # Загружаем каждый файл отдельно через ingest_text
    for filename in saved:
        filepath = os.path.join(base, filename)
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
                doc_id = f"file_{filename}_{uuid.uuid4().hex[:8]}"
                corpus.ingest_text(doc_id, content)
        except Exception as e:
            print(f"Error loading file {filename}: {e}")
            # Пробуем другие кодировки для файлов
            try:
                with open(filepath, 'r', encoding='cp1251') as f:
                    content = f.read()
                    doc_id = f"file_{filename}_{uuid.uuid4().hex[:8]}"
                    corpus.ingest_text(doc_id, content)
            except Exception as e2:
                print(f"Error loading file {filename} with cp1251: {e2}")
                continue
    
    append_trace({"type":"ingest_files", "files":saved})
    return {"ok": True, "files": saved, "count": len(saved)}

@app.post("/ingest")
async def ingest_unified(files: List[UploadFile] = File(None), text: str = Form(None), doc_id: Optional[str] = Form(None)):
    """Унифицированный эндпоинт для загрузки файлов и текста"""
    results = []
    
    # Обрабатываем файлы, если они есть
    if files:
        base = os.environ.get("DOCS_DIR", "data/docs"); os.makedirs(base, exist_ok=True)
        for f in files:
            fp = os.path.join(base, f.filename)
            with open(fp, "wb") as out: 
                out.write(await f.read())
            results.append({"type": "file", "filename": f.filename})
        
        # Загружаем каждый файл отдельно через ingest_text
        for f in files:
            try:
                filepath = os.path.join(base, f.filename)
                with open(filepath, 'r', encoding='utf-8') as file_content:
                    content = file_content.read()
                    doc_id = f"file_{f.filename}_{uuid.uuid4().hex[:8]}"
                    corpus.ingest_text(doc_id, content)
            except Exception as e:
                print(f"Error loading file {f.filename}: {e}")
                # Пробуем другие кодировки для файлов
                try:
                    with open(filepath, 'r', encoding='cp1251') as file_content:
                        content = file_content.read()
                        doc_id = f"file_{f.filename}_{uuid.uuid4().hex[:8]}"
                        corpus.ingest_text(doc_id, content)
                except Exception as e2:
                    print(f"Error loading file {f.filename} with cp1251: {e2}")
                    continue
    
    # Обрабатываем текст, если он есть
    if text and text.strip():
        if not doc_id:
            doc_id = f"text_{uuid.uuid4().hex[:8]}"
        
        try:
            import hashlib
            corpus.ingest_text(doc_id, text)
            results.append({"type": "text", "doc_id": doc_id})
        except Exception as e:
            import traceback
            print(f"Error during text ingestion: {e}")
            traceback.print_exc()
            return JSONResponse(status_code=500, content={"error": str(e)})
    
    if not results:
        return JSONResponse(status_code=400, content={"error": "No files or text provided"})
    
    append_trace({"type":"ingest_unified", "results": results})
    return {"ok": True, "results": results, "count": len(results)}

@app.post("/ingest/folder")
async def ingest_folder(path: str = Form(...)):
    corpus.ingest_folder(path)
    append_trace({"type":"ingest_folder", "path": path})
    return {"ok": True}

@app.get("/search")
def search(q: str, k: int = 5, entities: Optional[str] = None):
    # Фильтрация по сущностям здесь больше не поддерживается напрямую,
    # так как поиск теперь гибридный. Можно добавить в будущем.
    res = corpus.search(q, k=k)
    return JSONResponse(res)

@app.get("/ask")
async def ask(q: str, k: int = 5, entities: Optional[str] = None):
    ents = [x.strip() for x in entities.split(",")] if entities else None
    append_trace({"type":"query", "q": q, "entities": ents})
    res = await agent.run(q, k=k, entities_filter=ents)
    append_trace({"type":"result", "q": q, "answer": res.get("answer","")[:200], "citations": res.get("citations", [])})
    return JSONResponse(res)

@app.get("/ask/stream")
async def ask_stream(q: str, k: int = 5, entities: Optional[str] = None):
    """Эндпоинт для потоковой передачи RAG-ответа (классический MultiAgent)."""
    ents = [x.strip() for x in entities.split(",")] if entities else None
    
    async def event_generator():
        try:
            async for event in agent.stream(q, k=k, entities_filter=ents):
                yield {
                    "event": "message",
                    "data": json.dumps(event)
                }
        except Exception as e:
            # Отправляем событие с ошибкой на фронтенд
            error_event = {"type": "error", "data": f"Произошла ошибка: {e}"}
            yield {
                "event": "message",
                "data": json.dumps(error_event)
            }
            # Также логируем ошибку на сервере
            print(f"Error during stream: {e}")

    return EventSourceResponse(
        event_generator(),
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no"
        }
    )

@app.get("/ask/agentic")
async def ask_agentic(q: str, max_iterations: int = 5, confidence_threshold: float = 0.7):
    """🚀 Новый Agentic RAG эндпоинт - умная многоагентная обработка запросов."""
    try:
        result = await agentic_system.process_query(
            query=q, 
            max_iterations=max_iterations,
            confidence_threshold=confidence_threshold
        )
        append_trace({"type": "agentic_query", "q": q, "iterations": result.get("agentic_metadata", {}).get("iterations_used", 0)})
        return JSONResponse(result)
    except Exception as e:
        print(f"ERROR in /ask/agentic: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(status_code=500, content={"message": f"Agentic RAG error: {e}"})

@app.get("/ask/agentic/stream")
async def ask_agentic_stream(q: str, max_iterations: int = 5):
    """🚀 Потоковый Agentic RAG - показывает процесс принятия решений агентами."""
    
    async def agentic_event_generator():
        try:
            async for event in agentic_system.stream_query(q, max_iterations=max_iterations):
                yield {
                    "event": "message",
                    "data": json.dumps(event, ensure_ascii=False)
                }
        except Exception as e:
            error_event = {"type": "agentic_error", "data": f"Agentic RAG error: {e}"}
            yield {
                "event": "message", 
                "data": json.dumps(error_event)
            }
            print(f"Agentic RAG stream error: {e}")

    return EventSourceResponse(
        agentic_event_generator(),
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no"
        }
    )

@app.get("/documents")
def get_documents():
    """Возвращает список всех загруженных документов."""
    return JSONResponse(corpus.list_docs())

@app.get("/analytics/entities")
def get_entities_analytics():
    """Возвращает аналитику по сущностям в базе знаний."""
    try:
        # Получаем все сущности из графа
        all_entities = graph.get_all_entities_with_stats()
        
        # Подсчитываем статистику
        total_entities = len(all_entities)
        total_documents = len(corpus.list_docs())
        
        # Группируем по типам
        entity_types = {}
        entity_cloud = []
        
        for entity in all_entities:
            entity_type = entity.get('class', 'unknown')
            entity_text = entity.get('text', '')
            entity_count = entity.get('doc_count', 1)
            
            if entity_type not in entity_types:
                entity_types[entity_type] = 0
            entity_types[entity_type] += entity_count
            
            # Добавляем в облако с весом
            entity_cloud.append({
                'text': entity_text,
                'type': entity_type,
                'count': entity_count,
                'weight': min(entity_count * 10, 100)  # Вес для размера в облаке
            })
        
        # Находим топ категорию
        top_entity_type = max(entity_types.items(), key=lambda x: x[1])[0] if entity_types else 'Нет данных'
        
        # Сортируем облако по популярности
        entity_cloud.sort(key=lambda x: x['count'], reverse=True)
        entity_cloud = entity_cloud[:50]  # Топ-50 сущностей
        
        return JSONResponse({
            'total_entities': total_entities,
            'total_documents': total_documents,
            'top_entity_type': top_entity_type,
            'entity_types': entity_types,
            'entity_cloud': entity_cloud
        })
        
    except Exception as e:
        print(f"ERROR in /analytics/entities: {e}")
        return JSONResponse({
            'total_entities': 0,
            'total_documents': len(corpus.list_docs()) if corpus else 0,
            'top_entity_type': 'Ошибка',
            'entity_types': {},
            'entity_cloud': []
        })

@app.delete("/documents/{doc_id}")
def delete_document(doc_id: str):
    """Удаляет документ по его ID."""
    success = corpus.delete_doc(doc_id)
    if success:
        return JSONResponse({"status": "ok", "message": f"Document {doc_id} deleted."})
    else:
        return JSONResponse({"status": "error", "message": f"Document {doc_id} not found."}, status_code=404)

@app.post("/langextract/text")
async def langextract_text(task_prompt: str = Form(None), text: str = Form(...), doc_id: str = Form("extracted_doc")):
    try:
        out = run_extraction(text, prompt=task_prompt)
        graph.update_from_items(doc_id, out["items"])
        return JSONResponse(out)
    except Exception as e:
        print(f"ERROR in /langextract/text: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(status_code=500, content={"message": "Internal Server Error during extraction"})

@app.post("/langextract")
async def langextract_unified(request_data: dict = Body(...)):
    """Унифицированный эндпоинт для LangExtract с JSON входом"""
    try:
        text = request_data.get('text', '')
        task_prompt = request_data.get('task_prompt')
        doc_id = request_data.get('doc_id', 'extracted_doc')
        
        if not text.strip():
            return JSONResponse(status_code=400, content={"message": "Text is required", "success": False})
        
        out = run_extraction(text, prompt=task_prompt)
        graph.update_from_items(doc_id, out["items"])
        return JSONResponse(out)
    except Exception as e:
        print(f"ERROR in /langextract: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(status_code=500, content={"message": "Internal Server Error during extraction", "error": str(e), "success": False})

@app.get("/langextract/stream_text")
async def langextract_stream_text(text: str, task_prompt: Optional[str] = None, doc_id: str = "extracted_doc"):
    def gen():
        last = None
        for ev in stream_extraction(text, prompt=task_prompt):
            last = ev
            yield {"event": "message", "data": json.dumps(ev, ensure_ascii=False)}
        if last and isinstance(last, dict) and "result" in last:
            graph.update_from_items(doc_id, last["result"].get("items", []))
    return EventSourceResponse(gen())

@app.get("/extracts/{job_id}/viz.html")
async def get_viz(job_id: str):
    fp = os.path.join("data","extracts",job_id,"viz.html")
    if not os.path.exists(fp): return JSONResponse({"error":"not found"}, status_code=404)
    return FileResponse(fp, media_type="text/html")

@app.post("/api/extract-pdf-text")
async def extract_pdf_text(file: UploadFile = File(...)):
    """Извлечение текста из PDF файла"""
    try:
        if not file.filename.lower().endswith('.pdf'):
            return JSONResponse({"error": "Поддерживаются только PDF файлы"}, status_code=400)
        
        # Читаем содержимое файла
        content = await file.read()
        
        # Используем pypdf для извлечения текста
        try:
            from pypdf import PdfReader
            import io
            
            pdf_file = io.BytesIO(content)
            pdf_reader = PdfReader(pdf_file)
            
            text = ""
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text += f"\n--- Страница {page_num} ---\n{page_text}\n"
            
            if not text.strip():
                return JSONResponse({"error": "PDF файл не содержит текста или текст зашифрован"}, status_code=400)
            
            return JSONResponse({"text": text.strip()})
            
        except ImportError:
            # Fallback: используем pdfplumber если PyPDF2 недоступен
            try:
                import pdfplumber
                
                text = ""
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Страница {page_num} ---\n{page_text}\n"
                
                if not text.strip():
                    return JSONResponse({"error": "PDF файл не содержит текста"}, status_code=400)
                
                return JSONResponse({"text": text.strip()})
                
            except ImportError:
                return JSONResponse({
                    "error": "Для обработки PDF требуется установка PyPDF2 или pdfplumber. Скопируйте текст вручную."
                }, status_code=500)
        
    except Exception as e:
        print(f"Ошибка обработки PDF: {e}")
        return JSONResponse({"error": f"Ошибка обработки PDF: {str(e)}"}, status_code=500)

@app.post("/api/extract-text")
async def extract_text(file: UploadFile = File(...)):
    """Извлечение текста из различных форматов файлов"""
    try:
        filename = file.filename.lower()
        content = await file.read()
        
        # PDF файлы
        if filename.endswith('.pdf'):
            try:
                from pypdf import PdfReader
                import io
                
                pdf_file = io.BytesIO(content)
                pdf_reader = PdfReader(pdf_file)
                
                text = ""
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"\n--- Страница {page_num} ---\n{page_text}\n"
                
                return JSONResponse({"text": text.strip() or "PDF файл не содержит текста"})
                
            except Exception as e:
                return JSONResponse({"error": f"Ошибка обработки PDF: {str(e)}"}, status_code=500)
        
        # DOCX файлы
        elif filename.endswith('.docx'):
            try:
                import io
                from docx import Document
                
                print(f"Processing DOCX file: {filename}, size: {len(content)} bytes")
                doc_file = io.BytesIO(content)
                doc = Document(doc_file)
                
                text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text and paragraph.text.strip():
                        # Обеспечиваем правильную кодировку для русского текста
                        para_text = paragraph.text.strip()
                        text += para_text + "\n"
                
                # Дополнительно извлекаем текст из таблиц
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text and cell.text.strip():
                                cell_text = cell.text.strip()
                                text += cell_text + "\n"
                
                # Убеждаемся что текст в правильной кодировке
                if text.strip():
                    # Проверяем и исправляем возможные проблемы с кодировкой
                    try:
                        # Если текст содержит кракозябры, пробуем исправить
                        if '�' in text or any(ord(c) > 65535 for c in text):
                            print("Detected encoding issues in DOCX, attempting to fix...")
                            # Пробуем перекодировать
                            text = text.encode('utf-8', errors='ignore').decode('utf-8')
                    except Exception as encoding_error:
                        print(f"Encoding fix failed: {encoding_error}")
                
                return JSONResponse({
                    "text": text.strip() or "DOCX файл не содержит текста",
                    "format": "docx"
                })
                
            except ImportError:
                # Fallback: пробуем использовать docx2txt если python-docx недоступен
                try:
                    import docx2txt
                    import io
                    
                    # Сохраняем во временный файл для docx2txt
                    import tempfile
                    import os
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                        tmp_file.write(content)
                        tmp_file_path = tmp_file.name
                    
                    try:
                        text = docx2txt.process(tmp_file_path)
                        os.unlink(tmp_file_path)  # Удаляем временный файл
                        
                        return JSONResponse({
                            "text": text.strip() or "DOCX файл не содержит текста",
                            "format": "docx-fallback"
                        })
                    except Exception as docx2txt_error:
                        os.unlink(tmp_file_path)  # Удаляем временный файл
                        raise docx2txt_error
                        
                except ImportError:
                    return JSONResponse({
                        "error": "Для обработки DOCX требуется установка python-docx или docx2txt"
                    }, status_code=500)
            except Exception as e:
                import traceback
                error_details = traceback.format_exc()
                print(f"DOCX processing error: {e}")
                print(f"Full traceback: {error_details}")
                return JSONResponse({
                    "error": f"Ошибка обработки DOCX: {str(e)}",
                    "details": str(e)
                }, status_code=500)
        
        # Текстовые файлы с автоопределением кодировки
        elif filename.endswith(('.txt', '.md', '.rtf', '.csv')):
            try:
                import chardet
                
                # Определяем кодировку
                detected = chardet.detect(content)
                encoding = detected.get('encoding', 'utf-8')
                confidence = detected.get('confidence', 0)
                
                print(f"Detected encoding: {encoding} (confidence: {confidence})")
                
                # Пробуем декодировать с определенной кодировкой
                try:
                    text = content.decode(encoding)
                except UnicodeDecodeError:
                    # Fallback кодировки для русского языка
                    fallback_encodings = ['utf-8', 'windows-1251', 'cp1251', 'koi8-r']
                    text = None
                    
                    for enc in fallback_encodings:
                        try:
                            text = content.decode(enc)
                            print(f"Successfully decoded with fallback encoding: {enc}")
                            break
                        except UnicodeDecodeError:
                            continue
                    
                    if text is None:
                        text = content.decode('utf-8', errors='replace')
                        print("Used UTF-8 with error replacement")
                
                return JSONResponse({"text": text, "encoding": encoding, "confidence": confidence})
                
            except ImportError:
                # Если chardet не установлен, используем простые fallback'и
                fallback_encodings = ['utf-8', 'windows-1251', 'cp1251']
                for encoding in fallback_encodings:
                    try:
                        text = content.decode(encoding)
                        return JSONResponse({"text": text, "encoding": encoding})
                    except UnicodeDecodeError:
                        continue
                
                # Последний fallback
                text = content.decode('utf-8', errors='replace')
                return JSONResponse({"text": text, "encoding": "utf-8-replace"})
            
            except Exception as e:
                return JSONResponse({"error": f"Ошибка обработки текстового файла: {str(e)}"}, status_code=500)
        
        else:
            return JSONResponse({
                "error": f"Неподдерживаемый формат файла: {filename}"
            }, status_code=400)
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Общая ошибка извлечения текста: {e}")
        print(f"Full traceback: {error_details}")
        return JSONResponse({
            "error": f"Ошибка извлечения текста: {str(e)}",
            "type": "general_error",
            "filename": filename
        }, status_code=500)